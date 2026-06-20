"""Generate one synthetic clinical note in every supported upload format."""

from __future__ import annotations

import json
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "examples" / "cross_format"

TITLE = "Synthetic Clinical Review Note"
LINES = [
    "Patient: Amelia Hart",
    "Age: 67",
    "Sex: Female",
    "Document date: 2026-06-18",
    "Document type: Physician Progress Note",
    "",
    "Assessment",
    "Type 2 diabetes mellitus and hypertension.",
    "",
    "Symptoms",
    "Fatigue, blurred vision and dizziness. Patient denies chest pain.",
    "",
    "Current medications",
    "Metformin 1000 mg twice daily. Amlodipine 10 mg OD.",
    "",
    "Allergies",
    "Penicillin allergy.",
    "",
    "Results",
    "HbA1c 9.4%. Blood pressure 168/96 mmHg.",
    "",
    "Plan",
    "Route to the diabetes nurse specialist within 7 days.",
]
PLAIN_TEXT = "\n".join([TITLE, *LINES]).strip() + "\n"


def write_text_files() -> None:
    (OUTPUT_DIR / "clinical_review_note.txt").write_text(
        PLAIN_TEXT,
        encoding="utf-8",
    )
    markdown = "\n".join(
        [
            f"# {TITLE}",
            "",
            "**Patient:** Amelia Hart  ",
            "**Age:** 67  ",
            "**Sex:** Female  ",
            "**Document date:** 2026-06-18  ",
            "**Document type:** Physician Progress Note",
            "",
            "## Assessment",
            "Type 2 diabetes mellitus and hypertension.",
            "",
            "## Symptoms",
            "Fatigue, blurred vision and dizziness. Patient denies chest pain.",
            "",
            "## Current medications",
            "Metformin 1000 mg twice daily. Amlodipine 10 mg OD.",
            "",
            "## Allergies",
            "Penicillin allergy.",
            "",
            "## Results",
            "HbA1c 9.4%. Blood pressure 168/96 mmHg.",
            "",
            "## Plan",
            "Route to the diabetes nurse specialist within 7 days.",
            "",
        ]
    )
    (OUTPUT_DIR / "clinical_review_note.md").write_text(markdown, encoding="utf-8")


def write_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13)):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(46, 116, 181)
        style.font.bold = True

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(16)
    title_run = title.add_run(TITLE)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(23, 33, 43)

    for label, value in (
        ("Patient", "Amelia Hart"),
        ("Age", "67"),
        ("Sex", "Female"),
        ("Document date", "2026-06-18"),
        ("Document type", "Physician Progress Note"),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    sections = [
        ("Assessment", "Type 2 diabetes mellitus and hypertension."),
        (
            "Symptoms",
            "Fatigue, blurred vision and dizziness. Patient denies chest pain.",
        ),
        (
            "Current medications",
            "Metformin 1000 mg twice daily. Amlodipine 10 mg OD.",
        ),
        ("Allergies", "Penicillin allergy."),
        ("Results", "HbA1c 9.4%. Blood pressure 168/96 mmHg."),
        ("Plan", "Route to the diabetes nurse specialist within 7 days."),
    ]
    for heading, body in sections:
        document.add_heading(heading, level=2)
        document.add_paragraph(body)

    document.core_properties.title = TITLE
    document.core_properties.subject = "Synthetic cross-format POC input"
    document.core_properties.author = "Clinical Document Intelligence Hub"
    document.save(OUTPUT_DIR / "clinical_review_note.docx")


def write_pdf() -> None:
    path = OUTPUT_DIR / "clinical_review_note.pdf"
    pdf = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    pdf.setTitle(TITLE)
    pdf.setAuthor("Clinical Document Intelligence Hub")
    pdf.setFillColor(HexColor("#17212B"))
    pdf.setFont("Helvetica-Bold", 19)
    pdf.drawString(72, height - 72, TITLE)
    pdf.setStrokeColor(HexColor("#D8E0E5"))
    pdf.line(72, height - 84, width - 72, height - 84)

    y = height - 112
    headings = {"Assessment", "Symptoms", "Current medications", "Allergies", "Results", "Plan"}
    for line in LINES:
        if not line:
            y -= 7
            continue
        if line in headings:
            y -= 5
            pdf.setFillColor(HexColor("#2E74B5"))
            pdf.setFont("Helvetica-Bold", 12)
        else:
            pdf.setFillColor(HexColor("#17212B"))
            pdf.setFont("Helvetica", 10.5)
        pdf.drawString(72, y, line)
        y -= 17
    pdf.setFillColor(HexColor("#61707C"))
    pdf.setFont("Helvetica", 8)
    pdf.drawString(72, 42, "Synthetic demonstration record - no real patient data.")
    pdf.save()


def write_png() -> None:
    width, height = 1500, 1900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    regular_candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttf",
    ]
    bold_candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf",
    ]
    regular_path = next((path for path in regular_candidates if Path(path).exists()), None)
    bold_path = next((path for path in bold_candidates if Path(path).exists()), regular_path)
    regular = ImageFont.truetype(regular_path, 34) if regular_path else ImageFont.load_default()
    small = ImageFont.truetype(regular_path, 25) if regular_path else regular
    heading = ImageFont.truetype(bold_path, 39) if bold_path else regular
    title_font = ImageFont.truetype(bold_path, 54) if bold_path else regular

    x, y = 105, 100
    draw.text((x, y), TITLE, fill="#17212B", font=title_font)
    y += 88
    draw.line((x, y, width - x, y), fill="#D8E0E5", width=3)
    y += 44
    headings = {"Assessment", "Symptoms", "Current medications", "Allergies", "Results", "Plan"}
    for line in LINES:
        if not line:
            y += 20
            continue
        font = heading if line in headings else regular
        color = "#2E74B5" if line in headings else "#17212B"
        draw.text((x, y), line, fill=color, font=font)
        y += 58 if line in headings else 49

    draw.text(
        (x, height - 90),
        "Synthetic demonstration record - no real patient data.",
        fill="#61707C",
        font=small,
    )
    image.save(OUTPUT_DIR / "clinical_review_note.png", optimize=True)


def write_manifest() -> None:
    expected = {
        "patient": {"name": "Amelia Hart", "age": 67, "gender": "Female"},
        "document_date": "2026-06-18",
        "diagnoses": ["Type 2 diabetes mellitus", "hypertension"],
        "positive_symptoms": ["fatigue", "blurred vision", "dizziness"],
        "negated_symptoms": ["chest pain"],
        "medications": ["Metformin 1000 mg twice daily", "Amlodipine 10 mg OD"],
        "allergies": ["Penicillin allergy"],
        "results": ["HbA1c 9.4%", "Blood pressure 168/96 mmHg"],
        "expected_priority": "Medium",
    }
    (OUTPUT_DIR / "expected_output.json").write_text(
        json.dumps(expected, indent=2),
        encoding="utf-8",
    )


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    write_text_files()
    write_docx()
    write_pdf()
    write_png()
    write_manifest()
    print(f"Generated cross-format examples in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
